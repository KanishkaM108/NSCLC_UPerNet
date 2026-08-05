import os
import re
from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_formatted_text(paragraph, text):
    # Splits text by bold (**text**) and italic (*text*), formatting appropriately
    pattern = r'(\*\*.*?\*\*|\*.*?\*|`.*?`)'
    tokens = re.split(pattern, text)
    for token in tokens:
        if not token:
            continue
        if token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith('`') and token.endswith('`'):
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        else:
            paragraph.add_run(token)

def convert_md_to_docx(md_path, docx_path, img_dir):
    doc = docx.Document()

    # Set page margins (1 inch all around)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base font styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    num_lines = len(lines)
    
    in_table = False
    table_lines = []

    while i < num_lines:
        line = lines[i].rstrip('\r\n')

        # Check for markdown table start
        if '|' in line and i + 1 < num_lines and '|' in lines[i+1] and '---' in lines[i+1]:
            table_lines = [line]
            i += 1
            table_lines.append(lines[i].rstrip('\r\n')) # delimiter line
            i += 1
            while i < num_lines and '|' in lines[i]:
                table_lines.append(lines[i].rstrip('\r\n'))
                i += 1
            
            # Process table
            header_cells = [c.strip() for c in table_lines[0].strip('|').split('|')]
            rows_cells = [[c.strip() for c in r.strip('|').split('|')] for r in table_lines[2:]]
            
            table = doc.add_table(rows=len(rows_cells) + 1, cols=len(header_cells))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            set_table_borders(table)

            # Format Header
            hdr_cells = table.rows[0].cells
            for col_idx, h_text in enumerate(header_cells):
                hdr_cells[col_idx].text = h_text
                set_cell_background(hdr_cells[col_idx], "1A365D") # Navy blue header
                set_cell_margins(hdr_cells[col_idx], top=120, bottom=120, left=150, right=150)
                p = hdr_cells[col_idx].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    r.font.size = Pt(10)

            # Format Data Rows
            for r_idx, row_data in enumerate(rows_cells):
                row_cells = table.rows[r_idx + 1].cells
                bg_color = "F7FAFC" if r_idx % 2 == 1 else "FFFFFF"
                for col_idx, cell_value in enumerate(row_data):
                    if col_idx < len(row_cells):
                        p = row_cells[col_idx].paragraphs[0]
                        p.text = ""
                        add_formatted_text(p, cell_value)
                        set_cell_background(row_cells[col_idx], bg_color)
                        set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=150, right=150)
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.space_before = Pt(2)
                        for r in p.runs:
                            r.font.size = Pt(9.5)

            p_space = doc.add_paragraph()
            p_space.paragraph_format.space_after = Pt(6)
            continue

        # Check for Horizontal Rule
        if line.strip() in ['---', '***', '___']:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Check for Headings
        if line.startswith('# '):
            h = doc.add_heading(line[2:].strip(), level=1)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(8)
            for r in h.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(18)
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
            i += 1
            continue
        elif line.startswith('## '):
            h = doc.add_heading(line[3:].strip(), level=2)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(6)
            for r in h.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(14)
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
            i += 1
            continue
        elif line.startswith('### '):
            h = doc.add_heading(line[4:].strip(), level=3)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(4)
            for r in h.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(12)
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
            i += 1
            continue
        elif line.startswith('#### '):
            h = doc.add_heading(line[5:].strip(), level=4)
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after = Pt(4)
            for r in h.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(11)
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
            i += 1
            continue

        # Check for Images
        img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line.strip())
        if img_match:
            alt_text, img_src = img_match.groups()
            # Resolve image file
            img_file = None
            if img_src.startswith('file:///'):
                img_file = Path(img_src.replace('file:///', ''))
            elif img_src.startswith('/'):
                img_file = Path(img_src.lstrip('/'))
            else:
                img_file = img_dir / img_src

            if img_file and img_file.exists():
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(12)
                p_img.paragraph_format.space_after = Pt(4)
                run = p_img.add_run()
                run.add_picture(str(img_file), width=Inches(6.0))
            i += 1
            continue

        # Check for Figure/Table captions or Italic text
        if line.startswith('*Figure ') or line.startswith('*Table '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(12)
            add_formatted_text(p, line.strip())
            for r in p.runs:
                r.font.size = Pt(9.5)
                r.font.italic = True
                r.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
            i += 1
            continue

        # Check for Bullet Points
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(3)
            bullet_text = line.strip()[2:]
            add_formatted_text(p, bullet_text)
            i += 1
            continue

        # Check for Numbered Lists
        num_match = re.match(r'^(\d+)\.\s+(.*)', line.strip())
        if num_match:
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(3)
            add_formatted_text(p, num_match.group(2))
            i += 1
            continue

        # Regular Paragraph
        if line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            add_formatted_text(p, line.strip())
        
        i += 1

    doc.save(docx_path)
    print(f"Successfully generated DOCX at {docx_path}")

if __name__ == '__main__':
    md_file = Path("MANUSCRIPT_NATURE_SCIENTIFIC_REPORTS.md")
    docx_file = Path("MANUSCRIPT_NATURE_SCIENTIFIC_REPORTS_FINAL.docx")
    img_directory = Path(r"C:\Users\moham\.gemini\antigravity-ide\brain\d5f29b0a-6e40-4fc3-962b-2a4242ecb53e")
    convert_md_to_docx(md_file, docx_file, img_directory)
